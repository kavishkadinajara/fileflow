"""
Semantic Fidelity Index (SFI) — Gap 1 of the FileFlowOne research plan.

POST /api/slm-score
Measures how much semantic meaning survives when a document is converted
between formats using a three-dimensional weighted score.

  SFI = 0.35 · S_structural + 0.45 · S_semantic + 0.20 · S_functional
"""

from __future__ import annotations

import io
import re
import time
from typing import Any

from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from sentence_transformers import SentenceTransformer, util as st_util

_EMBED_MODEL: SentenceTransformer | None = None


def _get_embed_model() -> SentenceTransformer:
    global _EMBED_MODEL
    if _EMBED_MODEL is None:
        _EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
    return _EMBED_MODEL


router = APIRouter(prefix="/api", tags=["SFI"])

# ─────────────────────────────────────────────────────────────────────────────
# Format detection
# ─────────────────────────────────────────────────────────────────────────────

SUPPORTED_EXTS = {"docx", "pdf", "html", "md", "txt"}

def _detect_format(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext if ext in SUPPORTED_EXTS else "txt"


# ─────────────────────────────────────────────────────────────────────────────
# Extractors — return (plain_text, elements_dict)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> tuple[str, dict[str, int]]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    text_parts: list[str] = []
    headings = tables = lists = links = 0

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        text = para.text.strip()
        if text:
            text_parts.append(text)
        if "heading" in style_name:
            headings += 1
        if style_name in {"list bullet", "list number", "list paragraph"}:
            lists += 1

    tables = len(doc.tables)
    xml_str = doc.element.xml if hasattr(doc.element, "xml") else ""
    links = xml_str.count("<w:hyperlink ")

    return "\n".join(text_parts), {
        "headings": headings, "tables": tables,
        "lists": lists, "links": links,
    }


def _extract_pdf(data: bytes) -> tuple[str, dict[str, int]]:
    from pdfminer.high_level import extract_text
    from pdfminer.layout import LAParams

    params = LAParams(line_margin=0.5, word_margin=0.1)
    plain_text = extract_text(io.BytesIO(data), laparams=params) or ""

    lines = [l for l in plain_text.splitlines() if l.strip()]
    headings = lists = 0
    # Track table-like lines: require >= 3 multi-space gaps in a single line
    # AND the line must be fairly short (not a sentence). This avoids
    # over-counting ER diagram attribute lines as tables.
    table_lines = 0
    for line in lines:
        stripped = line.strip()
        # Heading: short, title-cased or all-caps, no sentence punctuation
        if (
            5 < len(stripped) < 70
            and not stripped.endswith((".", ",", ";", ":", "?"))
            and (stripped.istitle() or stripped.isupper())
        ):
            headings += 1
        # List: starts with bullet or number
        if re.match(r"^[\u2022\u2023\u25e6\-\*\+]\s|\d+[\.\)]\s", stripped):
            lists += 1
        # Table row: multiple tab/multi-space separated columns, NOT a sentence
        if len(re.findall(r"\s{2,}", stripped)) >= 3 and len(stripped) < 120:
            table_lines += 1

    # Each table has at least 2 rows (header + data) — divide and cap
    tables = min(table_lines // 2, 20)

    return plain_text, {
        "headings": headings, "tables": tables,
        "lists": lists, "links": 0,
    }


def _extract_html(data: bytes) -> tuple[str, dict[str, int]]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "html.parser")
    headings = sum(len(soup.find_all(f"h{i}")) for i in range(1, 7))
    tables   = len(soup.find_all("table"))
    lists    = len(soup.find_all(["ul", "ol"]))
    links    = len(soup.find_all("a", href=True))
    plain    = soup.get_text(separator="\n", strip=True)

    return plain, {"headings": headings, "tables": tables, "lists": lists, "links": links}


def _extract_md(data: bytes) -> tuple[str, dict[str, int]]:
    text = data.decode("utf-8", errors="replace")

    headings = len(re.findall(r"^#{1,6}\s", text, re.MULTILINE))
    # A table needs both a header row and a separator row (|---|)
    table_headers = re.findall(r"^\|.+\|$", text, re.MULTILINE)
    separators    = re.findall(r"^\|[\s\-:|]+\|$", text, re.MULTILINE)
    tables        = min(len(table_headers), len(separators))
    lists  = len(re.findall(r"^[\-\*\+]\s|\d+\.\s", text, re.MULTILINE))
    links  = len(re.findall(r"\[.+?\]\(.+?\)", text))
    formulas = len(re.findall(r"\$[^$\n]+\$|\\\(.+?\\\)", text))

    # Strip markdown syntax for plain text
    plain = re.sub(r"```[\s\S]*?```", " ", text)          # code blocks
    plain = re.sub(r"`[^`]+`", " ", plain)                 # inline code
    plain = re.sub(r"^#{1,6}\s+", "", plain, flags=re.MULTILINE)  # headings
    plain = re.sub(r"[*_~>|]", " ", plain)                 # formatting chars
    plain = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", plain) # keep link text
    plain = re.sub(r"\s+", " ", plain).strip()

    return plain, {
        "headings": headings, "tables": tables,
        "lists": lists, "links": links, "formulas": formulas,
    }


def _extract_txt(data: bytes) -> tuple[str, dict[str, int]]:
    text = data.decode("utf-8", errors="replace")
    paragraphs = [p for p in re.split(r"\n{2,}", text) if p.strip()]
    return text, {"headings": 0, "tables": 0, "lists": 0, "links": 0, "paragraphs": len(paragraphs)}


_EXTRACTORS = {
    "docx": _extract_docx, "pdf": _extract_pdf,
    "html": _extract_html, "md":  _extract_md, "txt": _extract_txt,
}


def extract(data: bytes, fmt: str) -> tuple[str, dict[str, int]]:
    return _EXTRACTORS.get(fmt, _extract_txt)(data)


# ─────────────────────────────────────────────────────────────────────────────
# Structural scoring
# ─────────────────────────────────────────────────────────────────────────────

def _ratio(target: int | float, source: int | float) -> float:
    if source == 0:
        return 1.0          # source has none → nothing to lose → full score
    return min(target / source, 1.0)


def score_structural(
    src: dict[str, int],
    tgt: dict[str, int],
    tgt_fmt: str,
) -> tuple[float, dict[str, Any]]:
    """
    Cross-format aware structural scoring.

    PDF output cannot preserve headings/lists/links in machine-readable form,
    so we do not penalize those dimensions when the target is PDF. Similarly,
    TXT loses all structure by definition.
    """
    lossy_target = tgt_fmt in {"pdf", "txt"}

    if lossy_target:
        # Structure is intentionally destroyed → give full structural score
        # but note it in details so the researcher knows.
        details = {
            "headings_source":    src.get("headings", 0),
            "headings_preserved": "N/A (lossy format)",
            "tables_source":      src.get("tables", 0),
            "tables_preserved":   "N/A (lossy format)",
            "lists_source":       src.get("lists", 0),
            "lists_preserved":    "N/A (lossy format)",
            "note": f"{tgt_fmt.upper()} does not preserve structural markup",
        }
        return 1.0, details

    heading_score = _ratio(tgt.get("headings", 0), src.get("headings", 0))
    table_score   = _ratio(tgt.get("tables",   0), src.get("tables",   0))
    list_score    = _ratio(tgt.get("lists",    0), src.get("lists",    0))

    score = heading_score * 0.4 + table_score * 0.4 + list_score * 0.2

    details = {
        "headings_source":    src.get("headings", 0),
        "headings_preserved": tgt.get("headings", 0),
        "tables_source":      src.get("tables", 0),
        "tables_preserved":   tgt.get("tables", 0),
        "lists_source":       src.get("lists", 0),
        "lists_preserved":    tgt.get("lists", 0),
        "heading_score":      round(heading_score, 4),
        "table_score":        round(table_score, 4),
        "list_score":         round(list_score, 4),
    }
    return round(score, 4), details


# ─────────────────────────────────────────────────────────────────────────────
# Semantic scoring
# ─────────────────────────────────────────────────────────────────────────────

def _split_chunks(text: str, max_words: int = 300) -> list[str]:
    """Split text into chunks of ~max_words at sentence boundaries."""
    # Split on sentence endings
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    chunks: list[str] = []
    current: list[str] = []
    count = 0
    for sent in sentences:
        words = sent.split()
        if not words:
            continue
        if count + len(words) > max_words and current:
            chunks.append(" ".join(current))
            current = list(words)
            count = len(words)
        else:
            current.extend(words)
            count += len(words)
    if current:
        chunks.append(" ".join(current))
    return [c for c in chunks if len(c.split()) >= 5]  # skip tiny fragments


def score_semantic(src_text: str, tgt_text: str) -> tuple[float, dict[str, Any]]:
    model = _get_embed_model()

    src_chunks = _split_chunks(src_text, max_words=300)
    tgt_chunks = _split_chunks(tgt_text, max_words=300)

    if not src_chunks or not tgt_chunks:
        # Fall back to full-document similarity if chunking yields nothing
        src_chunks = [src_text[:3000]] if src_text.strip() else []
        tgt_chunks = [tgt_text[:3000]] if tgt_text.strip() else []
        if not src_chunks or not tgt_chunks:
            return 0.0, {"chunks_analyzed": 0, "mean_similarity": 0.0,
                         "min_similarity": 0.0, "max_similarity": 0.0}

    src_emb = model.encode(src_chunks, convert_to_tensor=True, show_progress_bar=False)
    tgt_emb = model.encode(tgt_chunks, convert_to_tensor=True, show_progress_bar=False)

    cosine_scores = st_util.cos_sim(src_emb, tgt_emb)
    best_matches  = cosine_scores.max(dim=1).values

    mean_sim = float(best_matches.mean())
    min_sim  = float(best_matches.min())
    max_sim  = float(best_matches.max())

    return round(mean_sim, 4), {
        "chunks_analyzed": len(src_chunks),
        "mean_similarity": round(mean_sim, 4),
        "min_similarity":  round(min_sim, 4),
        "max_similarity":  round(max_sim, 4),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Functional scoring
# ─────────────────────────────────────────────────────────────────────────────

def _title_preserved(src_text: str, tgt_text: str, threshold: float = 0.6) -> bool:
    src_head  = src_text.strip()[:150].lower()
    tgt_lower = tgt_text.lower()
    src_words = set(re.findall(r"\w{4,}", src_head))  # only meaningful words
    tgt_words = set(re.findall(r"\w{4,}", tgt_lower[:800]))
    if not src_words:
        return True
    return len(src_words & tgt_words) / len(src_words) >= threshold


def score_functional(
    src: dict[str, int],
    tgt: dict[str, int],
    src_text: str,
    tgt_text: str,
    tgt_fmt: str,
) -> tuple[float, dict[str, Any]]:
    # Links: PDF/TXT cannot preserve clickable links → don't penalize
    if tgt_fmt in {"pdf", "txt"}:
        link_score = 1.0
    else:
        link_score = _ratio(tgt.get("links", 0), src.get("links", 0))

    metadata_score = 1.0 if _title_preserved(src_text, tgt_text) else 0.5

    # Formulas
    src_formulas = src.get("formulas", len(re.findall(r"\$[^$\n]+\$|\\\(.+?\\\)|=SUM", src_text)))
    tgt_formulas = tgt.get("formulas", len(re.findall(r"\$[^$\n]+\$|\\\(.+?\\\)|=SUM", tgt_text)))
    formula_score = _ratio(tgt_formulas, src_formulas)

    score = link_score * 0.5 + metadata_score * 0.3 + formula_score * 0.2

    return round(score, 4), {
        "link_score":         round(link_score, 4),
        "links_source":       src.get("links", 0),
        "links_preserved":    tgt.get("links", 0),
        "metadata_score":     round(metadata_score, 4),
        "formula_score":      round(formula_score, 4),
        "formulas_source":    src_formulas,
        "formulas_preserved": tgt_formulas,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Grade
# ─────────────────────────────────────────────────────────────────────────────

def _grade(sfi: float) -> str:
    if sfi >= 0.85: return "A"
    if sfi >= 0.70: return "B"
    if sfi >= 0.55: return "C"
    if sfi >= 0.40: return "D"
    return "F"


# ─────────────────────────────────────────────────────────────────────────────
# Route
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/slm-score")
async def slm_score(
    original_file:  UploadFile = File(..., description="Source document"),
    converted_file: UploadFile = File(..., description="Converted document"),
) -> JSONResponse:
    t_start = time.perf_counter()

    src_data = await original_file.read()
    tgt_data = await converted_file.read()

    src_fmt = _detect_format(original_file.filename or "file.txt")
    tgt_fmt = _detect_format(converted_file.filename or "file.txt")

    if not src_data or not tgt_data:
        raise HTTPException(status_code=422, detail="One or both files are empty.")

    try:
        src_text, src_elements = extract(src_data, src_fmt)
        tgt_text, tgt_elements = extract(tgt_data, tgt_fmt)
    except Exception as exc:
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=422,
            detail=f"Extraction failed ({src_fmt}→{tgt_fmt}): {type(exc).__name__}: {exc}",
        ) from exc

    if not src_text.strip():
        raise HTTPException(status_code=422, detail=f"Could not extract text from source file ({src_fmt}).")
    if not tgt_text.strip():
        raise HTTPException(status_code=422, detail=f"Could not extract text from converted file ({tgt_fmt}).")

    try:
        s_struct, struct_details = score_structural(src_elements, tgt_elements, tgt_fmt)
        s_sem,    sem_details    = score_semantic(src_text, tgt_text)
        s_func,   func_details   = score_functional(src_elements, tgt_elements, src_text, tgt_text, tgt_fmt)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Scoring failed: {exc}") from exc

    W_STRUCT, W_SEM, W_FUNC = 0.35, 0.45, 0.20
    sfi = round(W_STRUCT * s_struct + W_SEM * s_sem + W_FUNC * s_func, 4)

    elapsed_ms = round((time.perf_counter() - t_start) * 1000)

    return JSONResponse({
        "sfi_score": sfi,
        "grade":     _grade(sfi),
        "breakdown": {
            "structural": {
                "score":    s_struct,
                "weight":   W_STRUCT,
                "weighted": round(W_STRUCT * s_struct, 4),
                "details":  struct_details,
            },
            "semantic": {
                "score":    s_sem,
                "weight":   W_SEM,
                "weighted": round(W_SEM * s_sem, 4),
                "details":  sem_details,
            },
            "functional": {
                "score":    s_func,
                "weight":   W_FUNC,
                "weighted": round(W_FUNC * s_func, 4),
                "details":  func_details,
            },
        },
        "conversion_pair": {
            "source_format": src_fmt,
            "target_format": tgt_fmt,
        },
        "processing_time_ms": elapsed_ms,
    })
